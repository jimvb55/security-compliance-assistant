"""Document processing service for the Security Compliance Assistant."""
import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Iterator

from app.config import settings
from app.models.chunk import Chunk, ChunkMetadata
from app.models.document import Document, DocumentMetadata


class DocumentProcessor:
    """Service for processing documents (text extraction, chunking, etc.)."""
    
    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ):
        """Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks in words.
                If None, use the value from settings.
            chunk_overlap: Overlap between chunks in words.
                If None, use the value from settings.
        """
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
    
    def extract_text_from_file(self, file_path: Union[str, Path]) -> str:
        """Extract text from a file.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            Extracted text.
            
        Raises:
            ValueError: If the file type is not supported.
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif file_extension == ".docx":
            return self._extract_text_from_docx(file_path)
        elif file_extension == ".txt" or file_extension == ".md":
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Extracted text.
        """
        try:
            import PyPDF2
            
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                # Extract metadata if available
                metadata = reader.metadata
                if metadata:
                    if metadata.title:
                        text += f"Title: {metadata.title}\n"
                    if metadata.author:
                        text += f"Author: {metadata.author}\n"
                    if metadata.subject:
                        text += f"Subject: {metadata.subject}\n"
                    if metadata.creator:
                        text += f"Creator: {metadata.creator}\n"
                    text += "\n"
                
                # Extract text from each page
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                return text
        except ImportError:
            raise ImportError(
                "PyPDF2 is required to extract text from PDF files. "
                "Install it with 'pip install PyPDF2'."
            )
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Extracted text.
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract metadata if available
            core_properties = doc.core_properties
            if core_properties:
                if core_properties.title:
                    text += f"Title: {core_properties.title}\n"
                if core_properties.author:
                    text += f"Author: {core_properties.author}\n"
                if core_properties.subject:
                    text += f"Subject: {core_properties.subject}\n"
                if core_properties.comments:
                    text += f"Comments: {core_properties.comments}\n"
                text += "\n"
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text:
                        text += row_text + "\n"
            
            return text
        except ImportError:
            raise ImportError(
                "python-docx is required to extract text from DOCX files. "
                "Install it with 'pip install python-docx'."
            )
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from a text file.
        
        Args:
            file_path: Path to the text file.
            
        Returns:
            Extracted text.
        """
        with open(file_path, "r", encoding="utf-8", errors="replace") as file:
            return file.read()
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks of roughly equal size.
        
        Args:
            text: The text to chunk.
            
        Returns:
            List of text chunks.
        """
        # Split text into words
        words = re.findall(r'\S+|\n', text)
        
        if not words:
            return []
        
        # Create chunks
        chunks = []
        start = 0
        
        while start < len(words):
            # Calculate end of current chunk
            end = min(start + self.chunk_size, len(words))
            
            # Extract chunk
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap
        
        return chunks
    
    def create_chunks_from_document(self, document: Document) -> List[Chunk]:
        """Create chunks from a document.
        
        Args:
            document: The document to chunk.
            
        Returns:
            List of chunks.
        """
        # Split document text into chunks
        text_chunks = self.chunk_text(document.text)
        
        # Create Chunk objects
        chunks = []
        current_pos = 0
        
        for i, text_chunk in enumerate(text_chunks):
            # Calculate start and end positions
            start_idx = document.text.find(text_chunk, current_pos)
            if start_idx == -1:
                # If exact match not found, use approximate position
                start_idx = current_pos
            
            end_idx = start_idx + len(text_chunk)
            current_pos = end_idx
            
            # Create chunk metadata
            metadata = ChunkMetadata(
                doc_id=document.doc_id,
                chunk_id=f"{document.doc_id}_{i}",
                filename=document.metadata.filename,
                start_idx=start_idx,
                end_idx=end_idx,
                file_path=document.metadata.file_path,
                title=document.metadata.title,
                author=document.metadata.author,
                version=document.metadata.version,
                tags=document.metadata.tags.copy(),
                category=document.metadata.category,
                modified_date=document.metadata.modified_date.isoformat() if document.metadata.modified_date else None
            )
            
            # Create and add chunk
            chunk = Chunk(
                doc_id=document.doc_id,
                text=text_chunk,
                metadata=metadata
            )
            
            chunks.append(chunk)
        
        return chunks
    
    def extract_document_metadata(self, file_path: Union[str, Path]) -> DocumentMetadata:
        """Extract metadata from a document file.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            Document metadata.
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        # Create basic metadata from file path
        metadata = DocumentMetadata.from_file(file_path)
        
        # Try to extract additional metadata based on file type
        if file_extension == ".pdf":
            self._extract_pdf_metadata(file_path, metadata)
        elif file_extension == ".docx":
            self._extract_docx_metadata(file_path, metadata)
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path, metadata: DocumentMetadata) -> None:
        """Extract metadata from a PDF file.
        
        Args:
            file_path: Path to the PDF file.
            metadata: Metadata object to update.
        """
        try:
            import PyPDF2
            
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                pdf_metadata = reader.metadata
                
                if pdf_metadata:
                    if pdf_metadata.title:
                        metadata.title = pdf_metadata.title
                    if pdf_metadata.author:
                        metadata.author = pdf_metadata.author
                    if pdf_metadata.subject:
                        metadata.category = pdf_metadata.subject
        except (ImportError, Exception):
            # Silently fail if metadata extraction fails
            pass
    
    def _extract_docx_metadata(self, file_path: Path, metadata: DocumentMetadata) -> None:
        """Extract metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            metadata: Metadata object to update.
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            core_properties = doc.core_properties
            
            if core_properties:
                if core_properties.title:
                    metadata.title = core_properties.title
                if core_properties.author:
                    metadata.author = core_properties.author
                if core_properties.subject:
                    metadata.category = core_properties.subject
                if core_properties.version:
                    metadata.version = core_properties.version
        except (ImportError, Exception):
            # Silently fail if metadata extraction fails
            pass


# Default document processor instance
default_document_processor = DocumentProcessor()
