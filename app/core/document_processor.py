"""Document processing module for the Security Compliance Assistant."""
import io
import re
import uuid
from pathlib import Path
from typing import Dict, Generator, List, Optional, Tuple

import PyPDF2
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.config import settings
from app.models.document import Document, DocumentChunk, DocumentMetadata, DocumentType


class DocumentProcessor:
    """Processes documents for ingestion into the vector store."""
    
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        """Initialize the document processor.
        
        Args:
            chunk_size: Size of each chunk in words. Defaults to settings.CHUNK_SIZE.
            chunk_overlap: Overlap between chunks in words. Defaults to settings.CHUNK_OVERLAP.
        """
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
    
    def process_file(self, file_path: Path) -> Document:
        """Process a file and return a Document object.
        
        Args:
            file_path: Path to the file to process.
            
        Returns:
            Document object with metadata and chunks.
            
        Raises:
            ValueError: If the file type is not supported.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract metadata
        metadata = self._extract_metadata(file_path)
        
        # Extract text based on file type
        if metadata.doc_type == DocumentType.PDF:
            text, page_count = self._extract_text_from_pdf(file_path)
            metadata.page_count = page_count
        elif metadata.doc_type == DocumentType.DOCX:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx is required to process DOCX files")
            text, metadata = self._extract_text_from_docx(file_path, metadata)
        elif metadata.doc_type == DocumentType.TXT:
            text = self._extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {metadata.doc_type}")
        
        # Create document
        document = Document(metadata=metadata, raw_text=text)
        
        # Chunk the text
        document.chunks = self._chunk_text(text, document.metadata)
        
        # Update word count
        metadata.word_count = sum(len(chunk.text.split()) for chunk in document.chunks)
        
        return document
    
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract metadata from a file.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            DocumentMetadata object.
        """
        return DocumentMetadata.from_file(file_path)
    
    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Tuple of (text, page_count).
        """
        reader = PyPDF2.PdfReader(str(file_path))
        page_count = len(reader.pages)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, page_count
    
    def _extract_text_from_docx(self, file_path: Path, metadata: DocumentMetadata) -> Tuple[str, DocumentMetadata]:
        """Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            metadata: Document metadata to update.
            
        Returns:
            Tuple of (text, updated_metadata).
        """
        doc = docx.Document(file_path)
        
        # Extract metadata
        if doc.core_properties.title:
            metadata.title = doc.core_properties.title
        if doc.core_properties.author:
            metadata.author = doc.core_properties.author
        if doc.core_properties.created:
            metadata.created_date = doc.core_properties.created
        if doc.core_properties.modified:
            metadata.modified_date = doc.core_properties.modified
        
        # Extract text
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)
        
        # Update page count (approximation)
        metadata.page_count = len(doc.sections)
        
        return text, metadata
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file.
            
        Returns:
            Text content.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def _chunk_text(self, text: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Split text into chunks with overlap.
        
        Args:
            text: Text to chunk.
            metadata: Document metadata.
            
        Returns:
            List of DocumentChunk objects.
        """
        words = re.split(r"\s+", text)
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            if i > 0 and i + self.chunk_size > len(words):
                # If this is the last chunk and it would be smaller than chunk_size/2, 
                # merge it with the previous chunk
                if len(words) - i < self.chunk_size / 2:
                    break
            
            chunk_text = " ".join(words[i:i + self.chunk_size])
            if not chunk_text.strip():
                continue
                
            chunk_id = f"{metadata.doc_id}-{i//self.chunk_size}"
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                text=chunk_text,
                doc_id=metadata.doc_id,
                metadata=metadata,
                chunk_index=i // self.chunk_size
            )
            chunks.append(chunk)
            
        return chunks
